from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def _add_section_heading(doc: Document, text: str) -> None:
    """Add a Heading 2 styled as a labelled section divider."""
    doc.add_heading(text, level=2)


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    """Add each item as a bullet-list paragraph. Skips empty items."""
    for item in items:
        item = item.strip()
        if item:
            doc.add_paragraph(item, style="List Bullet")


def _add_none_if_empty(doc: Document, items: list[str]) -> None:
    """Add bullet list, or italicised 'None.' placeholder if empty."""
    if items:
        _add_bullet_list(doc, items)
    else:
        p = doc.add_paragraph()
        run = p.add_run("None.")
        run.italic = True


def build_docx(stories: list[dict], output_path: str) -> None:
    """
    Generate a formatted DOCX User Story Pack document.

    Each story renders all 14 sections:
      1. Epic & User Story Title
      2. User Story Statement
      3. Detailed Description
      4. Pre-Conditions
      5. Post-Conditions
      6. Data Governance
      7. Acceptance Criteria
      8. Assumptions
      9. Assertions (Out of Scope)
     10. Edge Cases & Error Scenarios
     11. Dependencies
     12. Example Data & Scenarios
     13. Test Scenarios
     14. Definition of Done
    """
    doc = Document()

    # Document title
    title_para = doc.add_heading("User Story Pack", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"{len(stories)} {'story' if len(stories) == 1 else 'stories'}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True

    for i, story in enumerate(stories):
        doc.add_paragraph()  # spacer

        # ── Section 1: Epic & Title ───────────────────────────────────────────
        epic = story.get("epic_title", "")
        if epic:
            epic_para = doc.add_paragraph()
            epic_run = epic_para.add_run(f"Epic: {epic}")
            epic_run.bold = True
            epic_run.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)  # indigo

        doc.add_heading(story.get("title", f"Story {i + 1}"), level=1)

        # Metadata line
        meta = doc.add_paragraph()
        meta.add_run(f"Priority: {story.get('priority', '?')}").bold = True
        meta.add_run(f"   |   Story Points: {story.get('story_points_estimate', '?')}")

        # ── Section 2: User Story Statement ──────────────────────────────────
        _add_section_heading(doc, "2. User Story Statement")
        statement = (
            f"As a {story.get('role', '')}, "
            f"I want {story.get('want', '')}, "
            f"so that {story.get('benefit', '')}."
        )
        doc.add_paragraph(statement)

        # ── Section 3: Detailed Description ──────────────────────────────────
        _add_section_heading(doc, "3. Detailed Description")
        desc = story.get("detailed_description", "").strip()
        doc.add_paragraph(desc if desc else "Not provided.")

        # ── Section 4: Pre-Conditions ─────────────────────────────────────────
        _add_section_heading(doc, "4. Pre-Conditions")
        _add_none_if_empty(doc, story.get("pre_conditions", []))

        # ── Section 5: Post-Conditions ────────────────────────────────────────
        _add_section_heading(doc, "5. Post-Conditions")
        _add_none_if_empty(doc, story.get("post_conditions", []))

        # ── Section 6: Data Governance ────────────────────────────────────────
        _add_section_heading(doc, "6. Data Governance")
        _add_none_if_empty(doc, story.get("data_governance", []))

        # ── Section 7: Acceptance Criteria ────────────────────────────────────
        _add_section_heading(doc, "7. Acceptance Criteria")
        _add_none_if_empty(doc, story.get("acceptance_criteria", []))

        # ── Section 8: Assumptions ────────────────────────────────────────────
        _add_section_heading(doc, "8. Assumptions")
        _add_none_if_empty(doc, story.get("assumptions", []))

        # ── Section 9: Assertions (Out of Scope) ──────────────────────────────
        _add_section_heading(doc, "9. Assertions (Limitations / Out of Scope)")
        _add_none_if_empty(doc, story.get("assertions", []))

        # ── Section 10: Edge Cases & Error Scenarios ──────────────────────────
        _add_section_heading(doc, "10. Edge Cases & Error Scenarios")
        _add_none_if_empty(doc, story.get("edge_cases", []))

        # ── Section 11: Dependencies ──────────────────────────────────────────
        _add_section_heading(doc, "11. Dependencies")
        _add_none_if_empty(doc, story.get("dependencies", []))

        # ── Section 12: Example Data & Scenarios ─────────────────────────────
        _add_section_heading(doc, "12. Example Data & Scenarios")
        _add_none_if_empty(doc, story.get("example_data", []))

        # ── Section 13: Test Scenarios ────────────────────────────────────────
        _add_section_heading(doc, "13. Test Scenarios")
        _add_none_if_empty(doc, story.get("test_scenarios", []))

        # ── Section 14: Definition of Done ───────────────────────────────────
        _add_section_heading(doc, "14. Definition of Done")
        _add_none_if_empty(doc, story.get("definition_of_done", []))

        # Story separator (except after last story)
        if i < len(stories) - 1:
            doc.add_paragraph()
            doc.add_paragraph("─" * 80)

    doc.save(output_path)
